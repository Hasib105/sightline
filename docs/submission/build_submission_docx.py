from __future__ import annotations

import base64
import json
import re
import sys
import textwrap
import time
import urllib.error
import urllib.request
import zlib
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / "docs" / "submission" / "project-documentation.md"
OUTPUT = ROOT / "docs" / "submission" / "Sightline_Project_Documentation.docx"
DIAGRAM_DIR = ROOT / "docs" / "submission" / "diagrams"

BASE_FONT = "Calibri"
HEADING_BLUE = "2E74B5"
HEADING_DARK = "1F4D78"
TABLE_HEADER_FILL = "F2F4F7"
BORDER = "D9E2EC"

MERMAID_INIT = """%%{init: {"theme": "base", "themeVariables": {
  "primaryColor": "#E8EEF5",
  "primaryTextColor": "#0B2545",
  "primaryBorderColor": "#2E74B5",
  "lineColor": "#4A5568",
  "secondaryColor": "#F8FAFC",
  "tertiaryColor": "#FFFFFF",
  "fontFamily": "Arial",
  "noteBkgColor": "#FFF7E6",
  "noteTextColor": "#1F2937"
}}}%%
"""


def slugify(value: str) -> str:
    value = re.sub(r"^\d+(\.\d+)*\.\s*", "", value.strip())
    value = re.sub(r"[^a-zA-Z0-9]+", "-", value).strip("-").lower()
    return value or "diagram"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int = 9360, indent_dxa: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_run_font(run, size: float | None = None, color: str | None = None) -> None:
    run.font.name = BASE_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), BASE_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BASE_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BASE_FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, HEADING_BLUE, 16, 8),
        ("Heading 2", 13, HEADING_BLUE, 12, 6),
        ("Heading 3", 12, HEADING_DARK, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = BASE_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BASE_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    title = styles["Title"]
    title.font.name = BASE_FONT
    title._element.rPr.rFonts.set(qn("w:eastAsia"), BASE_FONT)
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string("0B2545")
    title.paragraph_format.space_after = Pt(10)

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = BASE_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BASE_FONT)
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.10


def configure_document(doc: Document, landscape: bool = False) -> None:
    section = doc.sections[-1]
    section.top_margin = Inches(0.9 if landscape else 1.0)
    section.bottom_margin = Inches(0.8 if landscape else 1.0)
    section.left_margin = Inches(0.65 if landscape else 1.0)
    section.right_margin = Inches(0.65 if landscape else 1.0)
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)


def add_inline_markdown(paragraph, text: str, size: float | None = None) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            set_run_font(run, size=size)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            if size is not None:
                run.font.size = Pt(size)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size)


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    add_inline_markdown(p, text)


def parse_table_row(row: str) -> list[str]:
    row = row.strip()
    if row.startswith("|"):
        row = row[1:]
    if row.endswith("|"):
        row = row[:-1]
    return [cell.strip() for cell in row.split("|")]


def is_table_separator(row: str) -> bool:
    cells = parse_table_row(row)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def add_table(doc: Document, lines: list[str]) -> None:
    rows = [parse_table_row(line) for line in lines if line.strip()]
    if len(rows) >= 2 and is_table_separator(lines[1]):
        header = rows[0]
        body = rows[2:]
    else:
        header = rows[0]
        body = rows[1:]
    column_count = len(header)

    table = doc.add_table(rows=1, cols=column_count)
    table.style = "Table Grid"
    table.allow_autofit = False
    set_table_width(table)
    set_repeat_table_header(table.rows[0])

    widths = column_widths(header, body)
    for col_idx, text in enumerate(header):
        cell = table.rows[0].cells[col_idx]
        set_cell_width(cell, widths[col_idx])
        set_cell_margins(cell)
        set_cell_shading(cell, TABLE_HEADER_FILL)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text.replace("`", ""))
        run.bold = True
        set_run_font(run, size=8.6)

    for row in body:
        cells = table.add_row().cells
        for col_idx in range(column_count):
            value = row[col_idx] if col_idx < len(row) else ""
            cell = cells[col_idx]
            set_cell_width(cell, widths[col_idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            align_center = column_count >= 4 and len(value) <= 20
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align_center else WD_ALIGN_PARAGRAPH.LEFT
            add_inline_markdown(p, value, size=8.4 if column_count >= 4 else 8.8)

    doc.add_paragraph()


def column_widths(header: list[str], rows: list[list[str]]) -> list[int]:
    column_count = len(header)
    samples: list[int] = []
    for idx in range(column_count):
        values = [header[idx]]
        values.extend(row[idx] for row in rows if idx < len(row))
        avg = sum(min(len(value), 70) for value in values) / max(len(values), 1)
        samples.append(max(10, int(avg)))

    total = sum(samples) or column_count
    raw = [int(9360 * sample / total) for sample in samples]
    min_width = 1150 if column_count >= 5 else 1500 if column_count == 4 else 1800
    raw = [max(min_width, value) for value in raw]
    scale = 9360 / sum(raw)
    return [max(900, int(value * scale)) for value in raw]


def add_code_block(doc: Document, text: str) -> None:
    for line in text.rstrip("\n").splitlines() or [""]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.right_indent = Inches(0.15)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string("1F2937")


def kroki_png(source: str) -> bytes:
    payload = json.dumps({"diagram_source": source}).encode("utf-8")
    req = urllib.request.Request(
        "https://kroki.io/mermaid/png",
        data=payload,
        headers={
            "Content-Type": "application/json",
            "Accept": "image/png",
            "User-Agent": "Mozilla/5.0 SightlineDocumentationBuilder/1.0",
        },
        method="POST",
    )
    with urllib.request.urlopen(req, timeout=45) as response:
        return response.read()


def kroki_png_get(source: str) -> bytes:
    encoded = base64.urlsafe_b64encode(zlib.compress(source.encode("utf-8"), 9)).decode("ascii")
    req = urllib.request.Request(
        f"https://kroki.io/mermaid/png/{encoded}",
        headers={"User-Agent": "Mozilla/5.0 SightlineDocumentationBuilder/1.0"},
    )
    with urllib.request.urlopen(req, timeout=45) as response:
        return response.read()


def render_mermaid(source: str, name: str) -> Path:
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    output_path = DIAGRAM_DIR / f"{name}.png"
    if output_path.exists() and output_path.stat().st_size > 100:
        normalize_png_background(output_path)
        return output_path

    attempts = [MERMAID_INIT + source, source]
    last_error: Exception | None = None
    for attempt_source in attempts:
        for renderer in [kroki_png, kroki_png_get]:
            try:
                data = renderer(attempt_source)
                if not data.startswith(b"\x89PNG"):
                    raise RuntimeError("Renderer did not return PNG data.")
                output_path.write_bytes(data)
                normalize_png_background(output_path)
                time.sleep(0.2)
                return output_path
            except (urllib.error.URLError, urllib.error.HTTPError, TimeoutError, RuntimeError) as exc:
                last_error = exc
                time.sleep(0.5)
    raise RuntimeError(f"Could not render Mermaid diagram {name}: {last_error}")


def normalize_png_background(path: Path) -> None:
    image = Image.open(path)
    if image.mode in {"RGBA", "LA"} or ("transparency" in image.info):
        rgba = image.convert("RGBA")
        white = Image.new("RGBA", rgba.size, "WHITE")
        white.alpha_composite(rgba)
        white.convert("RGB").save(path)
    elif image.mode != "RGB":
        image.convert("RGB").save(path)


def add_sized_picture(paragraph, image_path: Path, max_width_in: float, max_height_in: float) -> None:
    image = Image.open(image_path)
    width_px, height_px = image.size
    aspect = width_px / max(height_px, 1)
    if max_width_in / aspect <= max_height_in:
        paragraph.add_run().add_picture(str(image_path), width=Inches(max_width_in))
    else:
        paragraph.add_run().add_picture(str(image_path), height=Inches(max_height_in))


def add_figure(doc: Document, image_path: Path, caption: str, source: str, figure_number: int) -> None:
    wide = source.lstrip().startswith("erDiagram")
    if wide:
        doc.add_section(WD_SECTION_START.NEW_PAGE)
        configure_document(doc, landscape=True)

    caption_text = f"Figure {figure_number}. {caption}"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run(caption_text)
    run.italic = True
    set_run_font(run, size=9.5, color="4A5568")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_sized_picture(p, image_path, max_width_in=9.3 if wide else 6.25, max_height_in=6.2 if wide else 8.0)

    if wide:
        doc.add_section(WD_SECTION_START.NEW_PAGE)
        configure_document(doc, landscape=False)


def strip_heading_number(text: str) -> str:
    return re.sub(r"^\d+(\.\d+)*\.\s*", "", text).strip()


def build_docx() -> None:
    doc = Document()
    configure_document(doc)
    configure_styles(doc)
    doc.core_properties.title = "Sightline Project Documentation"
    doc.core_properties.subject = "Submission documentation"
    doc.core_properties.author = "Sightline"

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    i = 0
    nearest_heading = "Diagram"
    figure_number = 1

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("```"):
            language = stripped.strip("`").strip().lower()
            block: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1
            block_text = "\n".join(block).strip("\n")
            if language == "mermaid":
                image_name = f"figure-{figure_number:02d}-{slugify(nearest_heading)}"
                image_path = render_mermaid(block_text, image_name)
                add_figure(doc, image_path, strip_heading_number(nearest_heading), block_text, figure_number)
                figure_number += 1
            else:
                add_code_block(doc, block_text)
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_table(doc, table_lines)
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading_match:
            hashes, heading_text = heading_match.groups()
            heading_text = heading_text.strip()
            nearest_heading = heading_text
            if len(hashes) == 1:
                p = doc.add_paragraph(style="Title")
                add_inline_markdown(p, heading_text)
                subtitle = doc.add_paragraph()
                subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = subtitle.add_run("Submission documentation with embedded project diagrams")
                set_run_font(run, size=11, color="4A5568")
            else:
                level = min(len(hashes) - 1, 3)
                doc.add_heading(heading_text, level=level)
            i += 1
            continue

        if re.match(r"^-\s+", stripped):
            while i < len(lines) and re.match(r"^-\s+", lines[i].strip()):
                item = re.sub(r"^-\s+", "", lines[i].strip())
                p = doc.add_paragraph(style="List Bullet")
                add_inline_markdown(p, item)
                i += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i].strip()):
                item = re.sub(r"^\d+\.\s+", "", lines[i].strip())
                p = doc.add_paragraph(style="List Number")
                add_inline_markdown(p, item)
                i += 1
            continue

        paragraph_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if (
                not nxt
                or nxt.startswith("#")
                or nxt.startswith("|")
                or nxt.startswith("```")
                or re.match(r"^-\s+", nxt)
                or re.match(r"^\d+\.\s+", nxt)
            ):
                break
            paragraph_lines.append(nxt)
            i += 1
        add_paragraph(doc, " ".join(paragraph_lines))

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    try:
        build_docx()
    except Exception as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        raise
